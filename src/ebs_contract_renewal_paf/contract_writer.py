"""
================================================================================
 Syntax Corporation © 2026 — All Rights Reserved
--------------------------------------------------------------------------------
 Project : EBS Contract Renewal PAF — BPA Renewal Automation Agent
 Module  : contract_writer.py — renewal contract document generator
 Version : 1.0.0      Build : 2026.06.04      Date : 2026-06-04
--------------------------------------------------------------------------------
 Produces the updated Blanket Purchase Agreement *document* (the blog's Step 3a)
 from the agent trace: parties, term, template legal language (governing law,
 warranty), and the revised pricing table (old vs new unit price + effective
 date). Companion to the PDOI CSV (Step 3b). Requires the optional `python-docx`
 package (pip install ".[docs]"); CSV/JSON output is unaffected if it is absent.
================================================================================
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

# Generic template clauses (a standard contract template, per the blog). These
# are a generic legal shell — customer legal owns the final language.
GOVERNING_LAW = ("This Agreement shall be governed by and construed in accordance "
                 "with the laws of the State of California, without regard to its "
                 "conflict-of-laws principles.")
WARRANTY = ("Supplier warrants that all goods supplied under this Agreement will "
            "be free from defects in material and workmanship for a period of "
            "twelve (12) months from delivery and will conform to the agreed "
            "specifications and quantities.")


def _hdr_run(p, text, size, color, *, bold=False, font="Georgia"):
    from docx.shared import Pt, RGBColor
    r = p.add_run(text)
    r.font.name = font; r.font.size = Pt(size); r.font.bold = bold
    r.font.color.rgb = RGBColor(*color)
    return r


def build_contract(trace: dict[str, Any], out_dir: str | Path,
                   assets_dir: str | Path | None = None) -> Path:
    """Render the renewal contract .docx from an agent trace dict."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:  # pragma: no cover - optional dep
        raise RuntimeError(
            "Contract document generation needs python-docx "
            "(pip install python-docx).") from exc

    NAVY = (0x06, 0x32, 0xA0); NAVY_DK = (0x04, 0x1F, 0x66)
    INK = (0x16, 0x20, 0x3A); SLATE = (0x5B, 0x65, 0x77)

    ag = trace.get("agreement_match") or {}
    sup = trace.get("supplier_match") or {}
    term = trace.get("term") or {}
    clog = trace.get("change_log") or []
    agreement_num = trace.get("agreement_num") or ag.get("agreement_num") or "—"
    eff = term.get("new_effective_date"); exp = term.get("new_expiration_date")

    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)
    st.font.color.rgb = RGBColor(*INK)

    if assets_dir and (Path(assets_dir) / "syntax-logo.png").exists():
        p = doc.add_paragraph()
        p.add_run().add_picture(str(Path(assets_dir) / "syntax-logo.png"),
                                width=Inches(1.7))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _hdr_run(p, "BLANKET PURCHASE AGREEMENT — RENEWAL", 20, NAVY_DK, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _hdr_run(p, f"Agreement No. {agreement_num}", 13, SLATE, font="Calibri")

    doc.add_paragraph()
    pr = doc.add_paragraph()
    pr.add_run("This Renewal is entered into between ")
    pr.add_run(sup.get("vendor_name") or "the Supplier").bold = True
    pr.add_run(" (“Supplier”) and the Buyer, renewing the above Blanket "
               "Purchase Agreement on the terms below.")

    # Term
    h = doc.add_paragraph(); _hdr_run(h, "1. Term", 13, NAVY, bold=True)
    doc.add_paragraph(f"The renewed term is effective {eff or '—'} through "
                      f"{exp or '—'} ({term.get('term_months', '—')} months), "
                      f"contiguous with the prior term ending "
                      f"{term.get('prior_expiration_date', '—')}.")

    # Pricing (old vs new) — the explainability change log
    h = doc.add_paragraph(); _hdr_run(h, "2. Revised pricing", 13, NAVY, bold=True)
    doc.add_paragraph("The unit prices for the renewed term are set out below, "
                      "reflecting current pricing retrieved from the system of "
                      "record with the agreed adjustment applied:")
    tbl = doc.add_table(rows=1, cols=6); tbl.style = "Table Grid"
    for i, hd in enumerate(["Line", "Item", "Prior $", "New $", "Δ", "Effective"]):
        c = tbl.rows[0].cells[i]
        run = c.paragraphs[0].add_run(hd); run.font.bold = True; run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
        sh.set(qn("w:fill"), "0632A0"); c._tc.get_or_add_tcPr().append(sh)
    for e in clog:
        cells = tbl.add_row().cells
        vals = [e.get("line_number"), e.get("item_number"),
                f"{e.get('old_unit_price'):.2f}" if e.get("old_unit_price") is not None else "—",
                f"{e.get('new_unit_price'):.2f}" if e.get("new_unit_price") is not None else "—",
                (f"{e.get('delta_pct')}%" if e.get("delta_pct") is not None else "—"),
                e.get("effective_date") or "—"]
        for i, v in enumerate(vals):
            r = cells[i].paragraphs[0].add_run(str(v)); r.font.size = Pt(10)

    # Governing law + warranty (template language)
    h = doc.add_paragraph(); _hdr_run(h, "3. Warranty", 13, NAVY, bold=True)
    doc.add_paragraph(WARRANTY)
    h = doc.add_paragraph(); _hdr_run(h, "4. Governing law", 13, NAVY, bold=True)
    doc.add_paragraph(GOVERNING_LAW)

    # Signatures
    h = doc.add_paragraph(); _hdr_run(h, "5. Signatures", 13, NAVY, bold=True)
    sg = doc.add_table(rows=4, cols=3); sg.style = "Table Grid"
    labels = ["Name", "Title", "Signature", "Date"]
    sg.rows[0].cells[1].paragraphs[0].add_run("Supplier").bold = True
    sg.rows[0].cells[2].paragraphs[0].add_run("Buyer").bold = True
    for i, lab in enumerate(labels[1:], 1):
        sg.rows[i].cells[0].paragraphs[0].add_run(lab).bold = True
    sg.rows[0].cells[0].paragraphs[0].add_run("Name").bold = True

    foot = doc.sections[0].footer.paragraphs[0]
    foot.text = "Generated by the EBS Contract Renewal PAF agent · Syntax Corporation © 2026 · Draft for review"
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in foot.runs:
        r.font.size = Pt(8); r.font.color.rgb = RGBColor(*SLATE)

    out = Path(out_dir); out.mkdir(parents=True, exist_ok=True)
    path = out / f"RENEWAL_CONTRACT_{str(agreement_num).replace('/', '-')}.docx"
    doc.save(path)
    return path
