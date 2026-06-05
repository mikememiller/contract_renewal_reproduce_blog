"""
================================================================================
 Syntax Corporation © 2026 — All Rights Reserved
--------------------------------------------------------------------------------
 Project : EBS Contract Renewal PAF — BPA Renewal Automation Agent
 Module  : deliverables/build_docx.py — install + tech-design + SOW (.docx)
 Version : 1.0.0      Build : 2026.06.04      Date : 2026-06-04
--------------------------------------------------------------------------------
 Generates three Syntax-branded Word documents into dist/ from the in-repo
 markdown + the one pricing model. Tech-design and SOW embed the architecture
 diagram; the SOW carries the 12/24/36-month pricing, a RACI, and a signature
 block (generic legal shell). All headings Georgia/navy, body Calibri.
================================================================================
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

import pricing

ROOT = Path(__file__).resolve().parents[1]
DIST = ROOT / "deliverables" / "dist"
ASSETS = ROOT / "deliverables" / "assets"
DIST.mkdir(parents=True, exist_ok=True)

NAVY = RGBColor(0x06, 0x32, 0xA0)
NAVY_DK = RGBColor(0x04, 0x1F, 0x66)
GREEN = RGBColor(0x3C, 0xC8, 0x5A)
SLATE = RGBColor(0x5B, 0x65, 0x77)
INK = RGBColor(0x16, 0x20, 0x3A)
HEAD = "Georgia"; BODY = "Calibri"
FOOTER = "Syntax Corporation © 2026 · Confidential"


def _base() -> Document:
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = BODY; st.font.size = Pt(11); st.font.color.rgb = INK
    for name, size in (("Heading 1", 17), ("Heading 2", 13.5), ("Heading 3", 12)):
        s = doc.styles[name]
        s.font.name = HEAD; s.font.bold = True; s.font.size = Pt(size)
        s.font.color.rgb = NAVY if name != "Heading 1" else NAVY_DK
    return doc


def _footer(doc):
    for sec in doc.sections:
        p = sec.footer.paragraphs[0]
        p.text = FOOTER; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = BODY; r.font.size = Pt(8); r.font.color.rgb = SLATE


def _cover(doc, kicker, title, subtitle):
    if (ASSETS / "syntax-logo.png").exists():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(str(ASSETS / "syntax-logo.png"), width=Inches(1.9))
    doc.add_paragraph()
    k = doc.add_paragraph(); r = k.add_run(kicker.upper())
    r.font.name = BODY; r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = NAVY
    t = doc.add_paragraph(); r = t.add_run(title)
    r.font.name = HEAD; r.font.bold = True; r.font.size = Pt(30); r.font.color.rgb = NAVY_DK
    s = doc.add_paragraph(); r = s.add_run(subtitle)
    r.font.name = BODY; r.font.size = Pt(13); r.font.color.rgb = SLATE
    d = doc.add_paragraph(); r = d.add_run("Syntax Corporation  ·  Build 1.0.0  ·  2026-06-04")
    r.font.name = BODY; r.font.size = Pt(10); r.font.color.rgb = SLATE
    doc.add_page_break()


def _shade(cell, hexcolor):
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor); cell._tc.get_or_add_tcPr().append(sh)


def _para(doc, text, size=11, color=INK, bold=False, italic=False):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name = BODY; r.font.size = Pt(size); r.font.color.rgb = color
    r.font.bold = bold; r.font.italic = italic
    return p


def _bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def _table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.font.name = BODY; run.font.bold = True; run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(c, "0632A0")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.name = BODY; run.font.size = Pt(10)
            if ri % 2 == 1:
                _shade(cells[i], "F4F7FC")
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


# ===========================================================================
def build_install():
    doc = _base()
    _cover(doc, "Installation Guide",
           "EBS Contract Renewal PAF", "Deployment & operation of the BPA renewal agent")
    doc.add_heading("1. Prerequisites", level=1)
    _bullets(doc, [
        "Python 3.10+ and (for the live path) Oracle Instant Client (thick mode — EBS enforces NNE).",
        "Network access to the EBS database (default host/SID via conn.json or env).",
        "Optional: ANTHROPIC_API_KEY for the LLM extractor; node for the sales-deck generator.",
    ])
    doc.add_heading("2. Install", level=1)
    _para(doc, "Run ./setup.sh — it creates a venv (--system-site-packages) and installs the "
               "agent, falling back to --no-deps on restricted networks.")
    doc.add_heading("3. Connect to EBS", level=1)
    _para(doc, "Settings precedence: flag > JSON (--config) > env > default. Copy conn.example.json "
               "to conn.json (gitignored); supply the password via EBS_PASSWORD or the interactive "
               "prompt — never hard-code it.")
    doc.add_heading("4. Run", level=1)
    _para(doc, "ebs-renewal-paf <quote.txt> --backend live --config conn.json", italic=True)
    _para(doc, "Outputs the two PDOI CSVs, agent_trace.json and qa_report.json into output/.")
    doc.add_heading("5. Load into EBS", level=1)
    _para(doc, "Recommended: hand the CSVs to the standard Import Price Catalogs (PDOI) program; "
               "the buyer reviews and approves the renewed agreement (it stages as INCOMPLETE). "
               "Optional direct staging via --load-to-ebs --yes (loadable batches only).")
    doc.add_heading("6. Deploy production PL/SQL (optional)", level=1)
    _bullets(doc, [
        "Deploy ebs/XXPAF_RENEWAL_PKG.pks/.pkb; run ebs/XXPAF_RENEWAL_utplsql.sql (parity).",
        "Register the managed-MCP tools (ebs/managed_mcp_tools.sql); grant the least-privilege read user.",
    ])
    _footer(doc)
    out = DIST / "EBS_Contract_Renewal_PAF_Installation_Guide.docx"
    doc.save(out); print(f"  OK  {out.name}")


def build_techdesign():
    doc = _base()
    _cover(doc, "Technical Design",
           "EBS Contract Renewal PAF", "Architecture, data contract & policy design")
    doc.add_heading("1. Overview", level=1)
    _para(doc, "A Private Agent Factory agent that renews EBS Blanket Purchase Agreements: it "
               "ingests a supplier renewal quote, pulls current prices from EBS via a governed MCP, "
               "applies a deterministic renewal-template policy, and stages a balanced Purchasing "
               "Documents Open Interface (PDOI) batch for the standard Import Price Catalogs program. "
               "Reproduces the Oracle blog 'Simplifying Contract Renewals: An AI Agent for EBS with "
               "Private Agent Factory'.")
    doc.add_heading("2. Reference architecture", level=1)
    if (DIST / "architecture.png").exists():
        doc.add_picture(str(DIST / "architecture.png"), width=Inches(6.4))
    _para(doc, "EBS 19c is the system of record (NNE; XXPAF PL/SQL). PAF runs on an ADB 26ai sidecar "
               "holding no EBS data and not in the data path. The OCI Database Tools Managed MCP "
               "(HTTPS + OAuth/IAM) reaches EBS via a Database Tools Connection. Writes (the load) "
               "stay off the MCP path.", italic=True)
    doc.add_heading("3. Pipeline", level=1)
    if (DIST / "process.png").exists():
        doc.add_picture(str(DIST / "process.png"), width=Inches(6.4))
    doc.add_heading("4. EBS interface contract (PDOI)", level=1)
    _para(doc, "Target: PO_HEADERS_INTERFACE + PO_LINES_INTERFACE; program: Import Price Catalogs. "
               "Verified live against EBS_Vision_12214:")
    _bullets(doc, [
        "ACTION='UPDATE' references the existing agreement by PO_HEADER_ID (DOCUMENT_NUM is an obsoleted stub).",
        "BATCH_ID is NUMBER; the agreement total is AMOUNT_AGREED (no BLANKET_TOTAL_AMOUNT on the interface).",
        "UOM_CODE for 'Each' is 'Ea'; PO_LINES_INTERFACE has no ATTRIBUTE columns.",
        "APPROVAL_STATUS='INCOMPLETE' — the buyer approves the renewal in EBS.",
    ])
    doc.add_heading("5. Renewal policy (deterministic)", level=1)
    _table(doc, ["Rule", "Logic", "Outcome"], [
        ["Price escalation", "increase > 7% AND annual $ impact > $250", "HOLD (dual gate)"],
        ["Term length", "renewed term > 36 months", "HOLD"],
        ["Term contiguity", "gap/overlap vs prior expiration", "HOLD"],
        ["Item validity", "item inactive / non-purchasable", "HOLD"],
        ["Supplier site", "site inactive / not a purchasing site", "HOLD"],
        ["Structural", "non-positive price / no matching line", "FAIL"],
        ["Balance", "AMOUNT_AGREED = Σ(price × qty)", "FAIL if unbalanced"],
    ], widths=[1.8, 3.4, 1.6])
    doc.add_heading("6. Quality & parity", level=1)
    _bullets(doc, [
        "43 hermetic + 8 live tests (golden record BPA 4467); read-only by default.",
        "Python reference engine is the test oracle; XXPAF_RENEWAL_PKG kept at parity (verified live).",
        "Deterministic policy outside the LLM → auditable, reproducible; full agent_trace.json.",
    ])
    _footer(doc)
    out = DIST / "EBS_Contract_Renewal_PAF_Technical_Design.docx"
    doc.save(out); print(f"  OK  {out.name}")


def build_sow():
    m = pricing.model()
    impl = m["implementation"]; msr = m["managed_services"]["terms"]; r = m["roi"]
    doc = _base()
    _cover(doc, "Statement of Work",
           "EBS Contract Renewal PAF", "Implementation & managed services — signable SOW")
    doc.add_heading("1. Scope", level=1)
    _para(doc, "Syntax Corporation will deliver a production Private Agent Factory agent that "
               "automates Oracle EBS Blanket Purchase Agreement renewals end-to-end: quote "
               "extraction, governed EBS reads, deterministic renewal policy, a balanced PDOI batch "
               "for Import Price Catalogs, a QA gate, the XXPAF PL/SQL package, and managed-MCP "
               "integration. Read-only by default; the buyer approves each renewal in EBS.")
    doc.add_heading("2. Architecture", level=1)
    if (DIST / "architecture.png").exists():
        doc.add_picture(str(DIST / "architecture.png"), width=Inches(6.2))
    doc.add_heading("3. Deliverables", level=1)
    _bullets(doc, [
        "The renewal agent (extractor, repository, policy engine, PDOI writer, QA gate, CLI).",
        "XXPAF_RENEWAL_PKG PL/SQL + parity harness; managed-MCP tool/SQL-Report definitions.",
        "Test suite (hermetic + live golden), QA report, and full documentation.",
        "PAF Agent Builder canvas recipe + published REST API hook.",
    ])
    doc.add_heading("4. Timeline", level=1)
    _table(doc, ["Phase", "Weeks", "Outcome"], [
        ["Discover & verify", "1–2", "Live SQL verification; golden record; spec sign-off"],
        ["Build vertical slice", "3–5", "Agent + PL/SQL + tests against the customer instance"],
        ["QA & UAT", "6–7", "Zero-bug convergence; balanced batch in UAT"],
        ["Cutover & enablement", "8", "Managed-MCP wiring; runbook; go-live"],
    ], widths=[2.4, 1.0, 3.2])
    _para(doc, f"Estimated effort: {impl['hours']} hours at a ${impl['blended_rate']:.0f}/hr blended "
               f"delivery rate ({int(pricing.DELIVERY_MIX['onshore']*100)}/"
               f"{int(pricing.DELIVERY_MIX['nearshore']*100)}/"
               f"{int(pricing.DELIVERY_MIX['offshore']*100)} onshore/nearshore/offshore).")
    doc.add_heading("5. RACI", level=1)
    _table(doc, ["Activity", "Syntax", "Customer"], [
        ["EBS discovery & SQL verification", "R/A", "C (access)"],
        ["Agent + PL/SQL build", "R/A", "C"],
        ["Tolerance / policy sign-off", "C", "R/A"],
        ["UAT & acceptance", "C", "R/A"],
        ["Production approval of renewals", "I", "R/A"],
        ["Managed services (run)", "R/A", "C"],
    ], widths=[3.0, 1.8, 1.8])
    doc.add_heading("6. Commercials", level=1)
    _para(doc, f"Fixed-fee implementation (one-time): ${impl['fixed_fee']:,.0f}.", bold=True)
    _table(doc, ["Managed services", "12-month", "24-month", "36-month"], [
        ["Per month", f"${msr['12']['monthly']:,.0f}", f"${msr['24']['monthly']:,.0f}",
         f"${msr['36']['monthly']:,.0f}"],
        ["Per year", f"${msr['12']['annual']:,.0f}", f"${msr['24']['annual']:,.0f}",
         f"${msr['36']['annual']:,.0f}"],
        ["Term discount", f"{msr['12']['discount_pct']}%", f"{msr['24']['discount_pct']}%",
         f"{msr['36']['discount_pct']}%"],
    ], widths=[2.4, 1.5, 1.5, 1.5])
    _para(doc, f"Indicative value: ${r['annual_savings']:,.0f} net annual savings, "
               f"~{r['payback_months']}-month payback, {r['three_year_roi_pct']:.0f}% 3-year ROI "
               f"(validate in a read-only value assessment). OCI consumption billed separately per "
               f"the OCI BOM.", italic=True)
    doc.add_heading("7. Acceptance criteria", level=1)
    _para(doc, "Green live golden-renewal test and a balanced PDOI batch staged and validated in the "
               "customer UAT environment.")
    doc.add_heading("8. Assumptions", level=1)
    _bullets(doc, [
        "Customer provides timely EBS access (read-only service account) and a test instance.",
        "Managed-MCP availability for the target EBS edition/region is confirmed at install.",
        "Legal terms are governed by the parties' master agreement; this SOW is a generic shell.",
    ])
    doc.add_heading("9. Signatures", level=1)
    _table(doc, ["", "Syntax Corporation", "Customer"], [
        ["Name", "", ""], ["Title", "", ""], ["Signature", "", ""], ["Date", "", ""],
    ], widths=[1.2, 2.8, 2.8])
    _footer(doc)
    out = DIST / "EBS_Contract_Renewal_PAF_SOW.docx"
    doc.save(out); print(f"  OK  {out.name}")


if __name__ == "__main__":
    build_install(); build_techdesign(); build_sow()
